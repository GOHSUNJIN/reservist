"""
Run:  python scripts/generate_doc.py
Output: ReservistGO_Overview.docx
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Palette ────────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x16, 0x1F, 0x30)
ACCENT = RGBColor(0x2F, 0x5F, 0xD0)
ACCENT_LIGHT = RGBColor(0x8A, 0xAA, 0xE8)
DARK   = RGBColor(0x22, 0x2D, 0x42)
MID    = RGBColor(0x4A, 0x5A, 0x72)
LIGHT  = RGBColor(0x8A, 0x94, 0xA3)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREEN  = RGBColor(0x1F, 0x8A, 0x5B)
AMBER  = RGBColor(0xB9, 0x79, 0x1A)
RED    = RGBColor(0xC0, 0x39, 0x2B)
GREEN_BG = 'E7F3EC'
AMBER_BG = 'FDF6E9'
RED_BG   = 'F7E4E1'

# ── XML helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def remove_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)

def set_table_full_width(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

def add_bookmark(para, bid, name):
    run = para.add_run()
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'),   str(bid))
    start.set(qn('w:name'), name)
    run._r.append(start)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bid))
    run._r.append(end)

def add_page_number_field(para):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '8A94A3')
    rPr.append(sz)
    rPr.append(color)
    r.append(rPr)
    fld1 = OxmlElement('w:fldChar')
    fld1.set(qn('w:fldCharType'), 'begin')
    r.append(fld1)
    para._p.append(r)
    r2 = OxmlElement('w:r')
    r2.append(rPr.copy() if hasattr(rPr, 'copy') else OxmlElement('w:rPr'))
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    r2.append(instr)
    para._p.append(r2)
    r3 = OxmlElement('w:r')
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'end')
    r3.append(fld2)
    para._p.append(r3)

def setup_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()

    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    '4')
    top.set(qn('w:space'), '4')
    top.set(qn('w:color'), 'C8D0DC')
    pBdr.append(top)
    pPr.append(pBdr)

    tabs_el = OxmlElement('w:tabs')
    t_right = OxmlElement('w:tab')
    t_right.set(qn('w:val'), 'right')
    t_right.set(qn('w:pos'), '9072')
    tabs_el.append(t_right)
    pPr.append(tabs_el)

    r_left = para.add_run('ReservistGO  |  Digital Attendance Management System')
    r_left.font.size  = Pt(8.5)
    r_left.font.color.rgb = LIGHT
    r_left.font.name  = 'Calibri'

    tab_r = OxmlElement('w:r')
    tab_e = OxmlElement('w:tab')
    tab_r.append(tab_e)
    para._p.append(tab_r)

    r_pg = para.add_run('Page ')
    r_pg.font.size  = Pt(8.5)
    r_pg.font.color.rgb = LIGHT
    r_pg.font.name  = 'Calibri'

    add_page_number_field(para)

# ── Layout helpers ─────────────────────────────────────────────────────────────

def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)
    return p

def page_break(doc):
    doc.add_page_break()

def section_banner(doc, number, title, bid=None, bname=None):
    """Full-width navy banner for H1 sections."""
    spacer(doc, 2)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(table)
    set_table_full_width(table)

    cell = table.rows[0].cells[0]
    set_cell_bg(cell, '161F30')
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(10)
    para.paragraph_format.left_indent  = Cm(0.5)

    n_run = para.add_run(f'{number:02d}   ')
    n_run.font.size  = Pt(11)
    n_run.font.bold  = True
    n_run.font.color.rgb = ACCENT_LIGHT
    n_run.font.name  = 'Calibri'

    t_run = para.add_run(title.upper())
    t_run.font.size  = Pt(15)
    t_run.font.bold  = True
    t_run.font.color.rgb = WHITE
    t_run.font.name  = 'Calibri'

    if bid is not None:
        add_bookmark(para, bid, bname)

    spacer(doc, 8)
    return table

def sub_heading(doc, title, bid=None, bname=None):
    """H2 with left accent border."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after  = Pt(5)
    para.paragraph_format.left_indent  = Cm(0.35)

    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '2F5FD0')
    pBdr.append(left)
    pPr.append(pBdr)

    run = para.add_run(title)
    run.font.size  = Pt(13)
    run.font.bold  = True
    run.font.color.rgb = ACCENT
    run.font.name  = 'Calibri'

    if bid is not None:
        add_bookmark(para, bid, bname)
    return para

def body(doc, text, lead=None):
    """Body paragraph with optional bold lead-in."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(5)
    if lead:
        r1 = para.add_run(lead + '  ')
        r1.bold = True
        r1.font.size  = Pt(11)
        r1.font.color.rgb = DARK
        r1.font.name  = 'Calibri'
    r2 = para.add_run(text)
    r2.font.size  = Pt(11)
    r2.font.color.rgb = MID
    r2.font.name  = 'Calibri'
    return para

def bullet(doc, text, label=None):
    """Bullet point with optional bold label."""
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(3)
    if label:
        r1 = para.add_run(label + ':  ')
        r1.bold = True
        r1.font.size  = Pt(11)
        r1.font.color.rgb = DARK
        r1.font.name  = 'Calibri'
    r2 = para.add_run(text)
    r2.font.size  = Pt(11)
    r2.font.color.rgb = MID
    r2.font.name  = 'Calibri'
    return para

def callout(doc, text, bg='EEF3FC', border='2F5FD0'):
    """Shaded callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(table)
    set_table_full_width(table)

    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'),   'single')
    left_b.set(qn('w:sz'),    '18')
    left_b.set(qn('w:space'), '0')
    left_b.set(qn('w:color'), border)
    tcBorders.append(left_b)
    tcPr.append(tcBorders)

    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(7)
    para.paragraph_format.space_after  = Pt(7)
    para.paragraph_format.left_indent  = Cm(0.3)

    run = para.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = MID
    run.font.name  = 'Calibri'
    run.font.italic = True

    spacer(doc, 6)
    return table

def styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(table)
    set_table_full_width(table)

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '2F5FD0')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after  = Pt(6)
        para.paragraph_format.left_indent  = Cm(0.2)
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = 'F4F7FC' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(5)
            para.paragraph_format.space_after  = Pt(5)
            para.paragraph_format.left_indent  = Cm(0.2)
            run = para.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = MID
            run.font.name = 'Calibri'

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    spacer(doc, 6)
    return table

def roadmap_table(doc, rows):
    effort_map = {
        'Quick win':  (GREEN_BG, '1F8A5B'),
        'Medium':     (AMBER_BG, 'B9791A'),
        'Large scope':(RED_BG,   'C0392B'),
    }
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(table)
    set_table_full_width(table)

    for i, h in enumerate(['#', 'Enhancement', 'Effort']):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, '161F30')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(7)
        para.paragraph_format.space_after  = Pt(7)
        para.paragraph_format.left_indent  = Cm(0.2)
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    for r_idx, (num, label, desc, effort) in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = 'F4F7FC' if r_idx % 2 == 0 else 'FFFFFF'
        ebg, efg = effort_map.get(effort, ('EEF3FC', '2F5FD0'))

        # Number cell
        set_cell_bg(row.cells[0], bg)
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_before = Pt(5)
        p0.paragraph_format.space_after  = Pt(5)
        p0.paragraph_format.left_indent  = Cm(0.2)
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(str(num))
        r0.bold = True
        r0.font.size = Pt(11)
        r0.font.color.rgb = ACCENT
        r0.font.name = 'Calibri'

        # Description cell
        set_cell_bg(row.cells[1], bg)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_before = Pt(5)
        p1.paragraph_format.space_after  = Pt(5)
        p1.paragraph_format.left_indent  = Cm(0.2)
        r1a = p1.add_run(label + '\n')
        r1a.bold = True
        r1a.font.size = Pt(10.5)
        r1a.font.color.rgb = DARK
        r1a.font.name = 'Calibri'
        r1b = p1.add_run(desc)
        r1b.font.size = Pt(9.5)
        r1b.font.color.rgb = MID
        r1b.font.name = 'Calibri'

        # Effort badge cell
        set_cell_bg(row.cells[2], ebg)
        p2 = row.cells[2].paragraphs[0]
        p2.paragraph_format.space_before = Pt(5)
        p2.paragraph_format.space_after  = Pt(5)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(effort)
        r2.bold = True
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = RGBColor(int(efg[0:2],16), int(efg[2:4],16), int(efg[4:6],16))
        r2.font.name = 'Calibri'

    for row in table.rows:
        row.cells[0].width = Cm(1.0)
        row.cells[1].width = Cm(12.2)
        row.cells[2].width = Cm(3.0)

    spacer(doc, 8)
    return table

def toc_entry(doc, num, label, bname, level=1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2 if level == 1 else 1)
    para.paragraph_format.space_after  = Pt(2 if level == 1 else 1)
    para.paragraph_format.left_indent  = Cm(0 if level == 1 else 1.0)

    if level == 1:
        r = para.add_run(f'{num}   {label}')
        r.font.size  = Pt(11.5)
        r.font.bold  = True
        r.font.color.rgb = DARK
        r.font.name  = 'Calibri'
    else:
        r = para.add_run(f'{num}  {label}')
        r.font.size  = Pt(10.5)
        r.font.bold  = False
        r.font.color.rgb = MID
        r.font.name  = 'Calibri'
    return para

# ── Build ──────────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)
        setup_footer(section)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    today = datetime.datetime.now().strftime('%B %Y')

    # ============================================================
    # COVER PAGE
    # ============================================================

    # Navy header block
    cover_table = doc.add_table(rows=1, cols=1)
    cover_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(cover_table)
    set_table_full_width(cover_table)

    cover_cell = cover_table.rows[0].cells[0]
    set_cell_bg(cover_cell, '161F30')
    cover_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cp = cover_cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(48)
    cp.paragraph_format.space_after  = Pt(6)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ct = cp.add_run('ReservistGO')
    ct.font.size  = Pt(42)
    ct.font.bold  = True
    ct.font.color.rgb = WHITE
    ct.font.name  = 'Calibri'

    cp2 = cover_cell.add_paragraph()
    cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp2.paragraph_format.space_before = Pt(0)
    cp2.paragraph_format.space_after  = Pt(48)
    cs = cp2.add_run('Digital Attendance Management System')
    cs.font.size  = Pt(16)
    cs.font.color.rgb = ACCENT_LIGHT
    cs.font.name  = 'Calibri'

    # Below the block
    for _ in range(4):
        spacer(doc, 14)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_r = meta_p.add_run(today.upper())
    meta_r.font.size  = Pt(11)
    meta_r.font.bold  = True
    meta_r.font.color.rgb = ACCENT
    meta_r.font.name  = 'Calibri'

    spacer(doc, 6)

    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_r = tag_p.add_run('Prepared for Supervisors and Operations Staff')
    tag_r.font.size   = Pt(11)
    tag_r.font.color.rgb = LIGHT
    tag_r.font.italic = True
    tag_r.font.name   = 'Calibri'

    spacer(doc, 10)

    # Thin accent rule
    rule_tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(rule_tbl)
    set_table_full_width(rule_tbl)
    set_cell_bg(rule_tbl.rows[0].cells[0], '2F5FD0')
    rule_tbl.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(2)
    rule_tbl.rows[0].cells[0].paragraphs[0].paragraph_format.space_after  = Pt(2)

    spacer(doc, 8)

    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_r = note_p.add_run(
        'This document provides an overview of the ReservistGO system, its features,\n'
        'and the planned enhancements roadmap. Intended for internal use only.'
    )
    note_r.font.size  = Pt(10)
    note_r.font.color.rgb = LIGHT
    note_r.font.name  = 'Calibri'

    # ============================================================
    # TABLE OF CONTENTS
    # ============================================================
    page_break(doc)

    toc_title = doc.add_paragraph()
    toc_title.paragraph_format.space_before = Pt(0)
    toc_title.paragraph_format.space_after  = Pt(4)
    tr = toc_title.add_run('Table of Contents')
    tr.font.size  = Pt(22)
    tr.font.bold  = True
    tr.font.color.rgb = NAVY
    tr.font.name  = 'Calibri'

    rule = doc.add_table(rows=1, cols=1)
    remove_table_borders(rule)
    set_table_full_width(rule)
    set_cell_bg(rule.rows[0].cells[0], '2F5FD0')
    rule.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(1)
    rule.rows[0].cells[0].paragraphs[0].paragraph_format.space_after  = Pt(1)

    spacer(doc, 10)

    toc_entries = [
        ('01', 'Overview',                        'sec_overview',  1),
        ('02', 'What Problem It Solves',           'sec_problem',   1),
        ('03', 'Access Roles',                     'sec_roles',     1),
        ('04', 'Features: For Reservists',         'sec_reservist', 1),
        ('',   '4.1  Check-In and GPS',            'sec_checkin',   2),
        ('',   '4.2  Leave and MC Requests',       'sec_leave',     2),
        ('',   '4.3  Attendance and Calendar',     'sec_attend',    2),
        ('',   '4.4  Safety and Reliability',      'sec_safety',    2),
        ('05', 'Features: For Supervisors',        'sec_admin',     1),
        ('',   '5.1  Attendance Roster and Log',   'sec_roster',    2),
        ('',   '5.2  Requests and Leave Inbox',    'sec_inbox',     2),
        ('',   '5.3  Personnel Management',        'sec_people',    2),
        ('',   '5.4  Cycle Management and Export', 'sec_cycle',     2),
        ('06', 'Features: For Master Level',       'sec_master',    1),
        ('07', 'Data Privacy',                     'sec_privacy',   1),
        ('08', 'Technical Overview',               'sec_tech',      1),
        ('09', 'Planned Enhancements',             'sec_roadmap',   1),
        ('',   '9.1  Quick Wins',                  'sec_r_quick',   2),
        ('',   '9.2  Medium Effort',               'sec_r_medium',  2),
        ('',   '9.3  Larger Scope',                'sec_r_large',   2),
    ]

    for num, label, bm, level in toc_entries:
        if level == 1 and num != toc_entries[0][0]:
            spacer(doc, 3)
        toc_entry(doc, num, label, bm, level)

    # ============================================================
    # SECTIONS 1-3: Overview, Problem, Roles  (same page group)
    # ============================================================
    page_break(doc)

    section_banner(doc, 1, 'Overview', bid=0, bname='sec_overview')
    body(doc,
        'ReservistGO is a mobile-first web application for managing NS reservist attendance. '
        'It replaces manual sign-in sheets and WhatsApp headcounts with a structured, auditable '
        'digital system accessible from any smartphone, with no app installation required.')
    body(doc,
        'Supervisors get a live dashboard of who has reported, who is late, and who has not shown up. '
        'Reservists check in from their phones using GPS verification. All records are timestamped, '
        'stored in the cloud, and exportable to Excel.')
    body(doc,
        'The system runs entirely in the browser. There is no hardware cost, no IT setup per user, '
        'and no dedicated terminals. Personnel access it through a URL that can be saved to the home '
        'screen for one-tap access.')

    callout(doc,
        'A built-in demo mode on the login screen allows anyone to explore both the reservist and '
        'supervisor views without creating an account. No registration is required to try the system.')

    # ── Section 2 ──────────────────────────────────────────────
    section_banner(doc, 2, 'What Problem It Solves', bid=1, bname='sec_problem')
    body(doc,
        'Before ReservistGO, attendance was tracked through a combination of paper sign-in sheets, '
        'manual WhatsApp headcounts, and verbal confirmation. This created several recurring problems:')
    spacer(doc, 4)

    bullet(doc, 'No single source of truth for who has reported. Supervisors had to consolidate information from multiple channels.')
    bullet(doc, 'No timestamped record. It was difficult to verify when a reservist actually checked in or whether they clocked out.')
    bullet(doc, 'Meal allowance errors. Without a reliable clock-out record, it was hard to confirm eligibility and handle corrections.')
    bullet(doc, 'Leave and MC tracking was informal. Requests came through messages and were not stored in a verifiable, auditable format.')
    bullet(doc, 'Absences were not automatically recorded. Someone had to manually identify and log no-shows each day.')

    spacer(doc, 6)
    body(doc, 'ReservistGO addresses all of these through a single system that all parties interact with directly.')

    # ── Section 3 ──────────────────────────────────────────────
    section_banner(doc, 3, 'Access Roles', bid=2, bname='sec_roles')
    body(doc, 'The system uses three permission levels. Each role sees only what is relevant to them.')
    spacer(doc, 4)

    styled_table(doc,
        ['Role', 'Who', 'Access'],
        [
            ('Reservist',  'NS personnel',              'Own check-in, leave requests, attendance history, and team directory. Scoped to their department.'),
            ('Supervisor', 'Staff officers, supervisors','Full roster, attendance log, time correction, leave approval, personnel records, cycle management, and exports. Scoped to their department.'),
            ('Master',     'Command level',             'Everything supervisors can do, plus managing all supervisor accounts and switching between departments.'),
        ],
        col_widths=[2.5, 3.8, 9.9]
    )

    # ============================================================
    # SECTION 4: Reservist Features
    # ============================================================
    page_break(doc)
    section_banner(doc, 4, 'Features: For Reservists', bid=3, bname='sec_reservist')
    body(doc,
        'All reservists access the system through the same URL. The interface adapts based on their department. '
        'Ops Security reservists see a four-phase check-in flow. CAS (Crime Alert) reservists see a '
        'simplified two-phase flow: check in and check out only.')

    sub_heading(doc, '4.1  Check-In and GPS', bid=4, bname='sec_checkin')
    bullet(doc, 'Ops Security reservists log four phases: check in, lunch out, return from lunch, and end of shift. CAS reservists log check in and check out only. Each phase is timestamped to the minute.', 'Department-based flow')
    bullet(doc, 'Tapping a phase first shows a "Locate me" button. Once GPS confirms the reservist is within range of HQ, the button changes to "Check in to work". Distance from HQ in metres is recorded against each entry. The accepted radius is configurable.', 'Sequential GPS verification')
    bullet(doc, 'After two failed GPS attempts, a bypass option appears. Bypassed records are permanently flagged in the log so supervisors are aware.', 'GPS bypass')
    bullet(doc, 'If the app is opened from WhatsApp, Instagram, or another in-app browser where GPS access is blocked, the app shows specific instructions for switching to the device browser.', 'In-app browser detection')
    bullet(doc, 'If connectivity is lost during check-in, the action is saved on the device and submitted automatically once the connection is restored.', 'Offline check-in')
    bullet(doc, 'If checking in more than one hour after shift start, the reservist is prompted for a written reason before the check-in is accepted. The reason is stored alongside the record.', 'Late declaration')

    sub_heading(doc, '4.2  Leave and MC Requests', bid=5, bname='sec_leave')
    bullet(doc, 'Reservists submit MC, personal leave, or other leave requests directly through the app. Requests go to the supervisor inbox with no phone calls or messages needed.', 'Digital submission')
    bullet(doc, 'A pending request can be withdrawn by the reservist before the supervisor has acted on it, from both the check-in screen and the Requests history.', 'Cancel pending requests')
    bullet(doc, 'The Requests history tab shows all past submissions with their current status: approved, rejected with reason, or withdrawn.', 'Request history')

    sub_heading(doc, '4.3  Attendance and Calendar', bid=6, bname='sec_attend')
    bullet(doc, 'A colour-coded calendar shows the full cycle record: present (green), MC (amber), absent, and no-report days (slate blue). Pending MC requests appear with a lighter tint and dashed border.', 'Attendance calendar')
    bullet(doc, 'If a past day shows no clock-out, that calendar cell is coloured orange with a dashed border. A persistent banner also appears on the check-in screen.', 'Missing clock-out detection')
    bullet(doc, 'A reminder banner appears automatically when a check-in phase window is open and the phase has not yet been recorded. It disappears once submitted.', 'Phase reminder banner')
    bullet(doc, 'All remaining no-report days in the current cycle are listed so reservists can plan ahead.', 'Upcoming no-report days')
    bullet(doc, 'When meal allowance is active for the cycle, a live timer shows total work time, paused during lunch. A "Meal eligible" badge appears once 6 hours of work is reached.', 'Work timer and meal eligibility')

    sub_heading(doc, '4.4  Safety and Reliability', bid=7, bname='sec_safety')
    bullet(doc, 'A confirmation step appears before logging out, preventing accidental session loss.', 'Logout confirmation')
    bullet(doc, 'A warning appears at 55 minutes with a one-tap option to extend the session. Sessions also expire after 20 minutes of inactivity, with a warning at 18 minutes.', 'Session and idle warnings')
    bullet(doc, 'Reservists can upload a profile photo from account settings. Tapping any photo in the app opens it enlarged.', 'Profile photo')

    # ============================================================
    # SECTION 5: Supervisor Features
    # ============================================================
    page_break(doc)
    section_banner(doc, 5, 'Features: For Supervisors', bid=8, bname='sec_admin')
    body(doc,
        'Supervisors have full visibility over their department\'s personnel, attendance records, leave requests, '
        'and cycle configuration. All data is scoped to the active department and updates in real time.')

    sub_heading(doc, '5.1  Attendance Roster and Log', bid=9, bname='sec_roster')
    bullet(doc, 'The roster updates in real time as reservists check in. No manual refresh is needed.', 'Live attendance board')
    bullet(doc, 'Filter by All, Present, MC, Absent, or Pending. Search by name. Navigate between dates by swipe or button, or jump directly to any date using the date picker.', 'Filters, search, navigation')
    bullet(doc, 'Mark any reservist as Present, MC, or Absent from the roster. Mark all pending as absent or all pending as present in a single action, each with a required confirmation step.', 'Status override')
    bullet(doc, 'Edit any reservist\'s check-in times for any phase on any day. Corrected records are flagged as admin-entered with the editor\'s name and timestamp.', 'Time correction')
    bullet(doc, 'Late check-ins (over one hour after shift start) are automatically flagged in the log with the reservist\'s written reason displayed alongside the record.', 'Late check-in alerts')
    bullet(doc, 'Any log entry where a reservist is marked present but has no clock-out is flagged with a visible badge, making it easy to identify before meal allowance is processed.', 'Missing clock-out flag')
    bullet(doc, 'Add a free-text note against any person\'s attendance entry. Also supports welfare notes and missed-attendance notes.', 'Log notes')

    sub_heading(doc, '5.2  Requests and Leave Inbox', bid=10, bname='sec_inbox')
    bullet(doc, 'All pending signup requests, MC requests, and leave requests appear in one unified inbox. Each signup is labelled New or Returning so supervisors know at a glance whether they are onboarding or re-enrolling.', 'Unified inbox')
    bullet(doc, 'Select multiple leave requests and approve or reject them in a single action. Bulk rejection requires a written reason stored against each rejected request.', 'Bulk leave actions')
    bullet(doc, 'Select-all checkbox in the signup section selects all visible pending signups, filtered by the search box, so it only acts on visible results.', 'Select-all for signups')
    bullet(doc, 'Rejected signup requests can be re-opened and returned to pending status if the decision needs to be reversed.', 'Reopen rejected signups')

    sub_heading(doc, '5.3  Personnel Management', bid=11, bname='sec_people')
    bullet(doc, 'Lists all personnel in the current cycle with their attendance rate. Personnel below 75% are flagged with an amber indicator.', 'Personnel roster')
    bullet(doc, 'Click any person\'s card to open their full attendance history across all cycles, with status filters and pagination. Exportable to Excel.', 'Per-person history')
    bullet(doc, 'Add multiple reservists at once by pasting a list of names and contact numbers.', 'Bulk add')
    bullet(doc, 'Reset any reservist\'s password directly from the roster without requiring database access.', 'Password reset')
    bullet(doc, 'Search across all cycles and all personnel by name, contact, or status. Supports permanent deletion and bulk delete.', 'Cross-cycle member search')

    sub_heading(doc, '5.4  Cycle Management and Export', bid=12, bname='sec_cycle')
    bullet(doc, 'Create and label reporting cycles. The system prepares the next 8 cycles automatically on every admin login.', 'Cycle management')
    bullet(doc, 'Mark any date as a no-report day. Paste a list of dates to mark multiple days at once. Singapore public holidays are excluded automatically.', 'No-report day control')
    bullet(doc, 'Post a short notice that appears on every reservist\'s check-in screen for the duration of the cycle. Scoped to the active department only.', 'Cycle notice board')
    bullet(doc, 'Enable or disable meal allowance per cycle. When active, the work timer and meal eligibility calculations are shown to reservists.', 'Meal allowance toggle')
    bullet(doc, 'Export the full attendance matrix as an Excel spreadsheet with per-person rates, colour-coded cells, and a meal claims column.', 'Excel export')
    bullet(doc, 'Generate a formatted A4 attendance report, printable or saveable as PDF directly from the browser.', 'Print report')
    bullet(doc, 'One tap generates the day\'s attendance summary as a WhatsApp-ready message, with a preview, copy, and direct-send option.', 'WhatsApp summary')

    # ============================================================
    # SECTION 6: Master Level  +  SECTION 7: Privacy  (same page)
    # ============================================================
    page_break(doc)
    section_banner(doc, 6, 'Features: For Master Level', bid=13, bname='sec_master')
    body(doc, 'The Master account holds all supervisor capabilities and the following additional controls:')
    spacer(doc, 4)

    bullet(doc, 'Switch the admin view between departments (Ops Security and Crime Alert/CAS) using a dropdown in the header. Each department\'s last-selected cycle is remembered independently.', 'Department switcher')
    bullet(doc, 'All data is fully scoped to the active department. Switching departments clears the current view and reloads fresh data for the selected department.', 'Cross-department isolation')
    bullet(doc, 'Add new supervisor accounts directly within the app. Demote any admin back to reservist status, or promote any reservist to admin.', 'Supervisor account management')
    bullet(doc, 'Reset any supervisor\'s password directly from the Team tab, without requiring database access.', 'Supervisor password reset')

    section_banner(doc, 7, 'Data Privacy', bid=14, bname='sec_privacy')
    body(doc, 'ReservistGO collects only what is operationally necessary. The table below summarises what is and is not held in the system.')
    spacer(doc, 6)

    styled_table(doc,
        ['What is stored', 'What is NOT stored'],
        [
            ('Name, phone number, attendance records',          'NRIC, rank, or service details'),
            ('Distance from HQ in metres at check-in only',     'Location history or exact GPS coordinates'),
            ('Profile photo (optional, removable at any time)', 'Device identifiers or browser fingerprints'),
            ('Encrypted passwords (via authentication service)','Passwords in plain text (inaccessible even to admins)'),
            ('Temporary browser session',                       'Persistent login data stored on the device'),
        ],
        col_widths=[8.0, 8.2]
    )

    callout(doc,
        'Sessions are held in the browser\'s temporary memory and are cleared when the tab is closed or after '
        '20 minutes of inactivity. Passwords are encrypted before storage. Not even the Master account can read '
        'a user\'s password.')

    # ============================================================
    # SECTION 8: Technical Overview
    # ============================================================
    page_break(doc)
    section_banner(doc, 8, 'Technical Overview', bid=15, bname='sec_tech')
    body(doc,
        'This section is a brief non-technical summary of the components that power the system. '
        'No specialist knowledge is needed to operate the system day-to-day.')
    spacer(doc, 6)

    styled_table(doc,
        ['Component', 'Technology', 'What it does'],
        [
            ('The app',        'Vanilla JavaScript',   'Runs entirely in the browser. No third-party framework or app installation required.'),
            ('Database',       'Supabase / PostgreSQL', 'Stores all personnel, attendance, and leave records. Managed, hosted, and automatically backed up.'),
            ('Login',          'Supabase Auth',         'Handles all password security. Passwords are encrypted before storage.'),
            ('Live updates',   'Supabase Realtime',     'Pushes attendance changes to all connected supervisors instantly, without any page refresh.'),
            ('Profile photos', 'Supabase Storage',      'Profile pictures stored in the cloud, isolated per user.'),
            ('Hosting',        'Vercel',                'Deployed globally on a CDN. Loads quickly regardless of network conditions.'),
            ('Offline support','Service Worker',        'Caches the app and queues check-in actions when there is no internet connection.'),
        ],
        col_widths=[3.0, 3.5, 9.7]
    )

    callout(doc,
        'The system requires no dedicated servers, no IT maintenance, and no per-device setup. '
        'Updates are deployed instantly to all users without anyone needing to refresh or reinstall.')

    # ============================================================
    # SECTION 9: Roadmap
    # ============================================================
    page_break(doc)
    section_banner(doc, 9, 'Planned Enhancements', bid=16, bname='sec_roadmap')
    body(doc,
        'The following improvements are planned to make the system fully self-managed by supervisors, '
        'requiring no developer involvement after initial deployment. '
        'Items are grouped by implementation effort and ordered from lowest to highest complexity.')

    spacer(doc, 4)
    styled_table(doc,
        ['Tier', 'Items', 'Description'],
        [
            ('Quick win',   '1-5',  'Configuration and database changes that can be delivered quickly with high operational impact.'),
            ('Medium',      '6-10', 'More involved changes requiring schema updates or new infrastructure components.'),
            ('Large scope', '11-14','Significant changes that make the system entirely self-contained after initial deployment.'),
        ],
        col_widths=[2.8, 1.5, 11.9]
    )

    sub_heading(doc, '9.1  Quick Wins', bid=17, bname='sec_r_quick')
    body(doc, 'Low-complexity changes with immediate operational value. Each can be delivered independently.')
    spacer(doc, 6)

    roadmap_table(doc, [
        (1, 'Editable reporting timings',
         'Phase windows (0900, 1200, 1400, 1800) are currently fixed constants. Moving them into the database per cycle lets supervisors set different reporting hours from within the app. Phase reminders, the work timer, and meal eligibility all update automatically.',
         'Quick win'),
        (2, 'Editable Info tab content',
         'Attire requirements, meal form links, and dekit checklists are hardcoded. An edit form in the cycle management panel lets supervisors keep this content current without touching any files.',
         'Quick win'),
        (3, 'WhatsApp group link in database',
         'The unit WA group link is currently set once at deployment. Moving it to a department record lets supervisors update it from within the app when the group changes.',
         'Quick win'),
        (4, 'Multiple shifts per department',
         'The system currently supports one shift type. Extending to AM, PM, Night, and custom shifts (each with configurable phase windows) supports departments with rotating schedules.',
         'Quick win'),
        (5, 'Per-cycle GPS location and radius',
         'HQ coordinates and the check-in radius are currently deployment-level settings. Storing them per cycle allows supervisors to run cycles at different venues (exercises, off-site operations) without a code deployment.',
         'Quick win'),
    ])

    sub_heading(doc, '9.2  Medium Effort', bid=18, bname='sec_r_medium')
    body(doc, 'More involved changes requiring database schema updates or new service integrations.')
    spacer(doc, 6)

    roadmap_table(doc, [
        (6, 'Department CRUD',
         'Departments are currently a fixed database type. Migrating to a departments table lets the Master account create and manage departments from within the app. This is the prerequisite for item 7.',
         'Medium'),
        (7, 'Per-department configuration panel',
         'Once departments are table-driven, each department stores its own HQ coordinates, phase windows, WA group link, and Info tab content. A Settings panel in the app replaces all deployment-time configuration.',
         'Medium'),
        (8, 'Shift scheduler',
         'Pre-assign which reservists report on which days. This generates the expected attendance list per date, so mark-all-absent only targets scheduled personnel rather than the full roster.',
         'Medium'),
        (9, 'Push notifications',
         'The service worker is already in place. Adding Web Push lets reservists receive reminders when a phase window opens, and supervisors receive alerts on new requests without needing the app open.',
         'Medium'),
        (10, 'Audit log viewer',
         'A read-only log of all supervisor actions (status overrides, time corrections, approvals, account changes) surfaced as a tab for the Master account for accountability and oversight.',
         'Medium'),
    ])

    sub_heading(doc, '9.3  Larger Scope', bid=19, bname='sec_r_large')
    body(doc, 'Significant changes that would make the system entirely self-contained after initial deployment, with no developer involvement required.')
    spacer(doc, 6)

    roadmap_table(doc, [
        (11, 'In-app configuration editor',
         'Replace all deployment-time configuration (org name, accent colour, HQ location) with a Settings panel inside the app. After initial deployment, no files ever need to be edited again.',
         'Large scope'),
        (12, 'Scoped department access for supervisors',
         'Currently all admins in a department see all data for that department. As more departments are added, supervisors may need to be scoped to one department, with the Master retaining full cross-department visibility.',
         'Large scope'),
        (13, 'Automated attendance summaries',
         'Scheduled daily and weekly reports sent to the supervisor\'s WhatsApp or email, covering attendance rate, pending leave requests, and missing clock-outs. Requires a messaging API integration.',
         'Large scope'),
        (14, 'Equipment and dekit tracking',
         'Track issued items per reservist per cycle and record return status at dekit. Fits naturally into the existing cycle lifecycle alongside the dekit date already stored on each cycle.',
         'Large scope'),
    ])

    # ── Save ───────────────────────────────────────────────────
    out = 'ReservistGO_Overview.docx'
    doc.save(out)
    print(f'Saved: {out}')

if __name__ == '__main__':
    build()
